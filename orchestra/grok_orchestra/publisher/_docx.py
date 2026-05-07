"""DOCX builder — split out so the heavy ``docx`` import stays lazy.

Word's built-in styles (`Heading 1` … `Normal`) are preserved so the
TOC field works without any post-processing in Word.
"""

from __future__ import annotations

from collections.abc import Mapping, Sequence
from pathlib import Path
from typing import Any

__all__ = ["write_docx"]


def write_docx(ctx: Mapping[str, Any], output_path: Path) -> Path:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Title (Word style "Title").
    title = doc.add_paragraph()
    run = title.add_run(str(ctx.get("title", "Grok Agent Orchestra Report")))
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0B, 0x12, 0x20)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"run · {ctx['run_id']}   |   {ctx['generated_at']}"
        + (f"   |   {ctx['template_name']}" if ctx.get("template_name") else "")
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    confidence = ctx["confidence"]
    verdict = doc.add_paragraph()
    badge = "✅ approved" if confidence.approved else "⛔ blocked"
    vrun = verdict.add_run(
        f"Lucas verdict: {badge}   ·   confidence "
        f"{int(round(confidence.overall * 100))}%"
    )
    vrun.bold = True
    vrun.font.size = Pt(11)
    vrun.font.color.rgb = (
        RGBColor(0x16, 0xA3, 0x4A) if confidence.approved else RGBColor(0xDC, 0x26, 0x26)
    )

    doc.add_paragraph()  # spacer

    image_refs = dict(ctx.get("image_refs") or {})
    image_dir = ctx.get("image_dir") or ""
    _embed_image(doc, image_refs.get("cover"), image_dir, width_in=6.0)

    _section(doc, "Executive Summary", ctx.get("executive_summary"))
    _section(doc, "Findings", ctx.get("findings"), default="Harper produced no per-role output for this run.", image=image_refs.get("findings"), image_dir=image_dir)
    _section(doc, "Analysis", ctx.get("analysis"), default="Benjamin produced no per-role output for this run.", image=image_refs.get("analysis"), image_dir=image_dir)
    _section(doc, "Stress Test", ctx.get("stress_test"), default="Lucas produced no per-role narrative — see verdict below.", image=image_refs.get("stress_test"), image_dir=image_dir)
    _section(doc, "Synthesis", ctx.get("synthesis") or ctx.get("executive_summary"), image=image_refs.get("synthesis"), image_dir=image_dir)

    # Lucas verdict table.
    doc.add_heading("Lucas Verdict", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    _row(table, "Approved", "yes" if confidence.approved else "no")
    _row(table, "Overall confidence", f"{confidence.overall:.2f}")
    if confidence.per_claim:
        _row(
            table,
            "Per-claim scores",
            ", ".join(f"{k}: {v:.2f}" for k, v in confidence.per_claim.items()),
        )
    if confidence.warnings:
        _row(table, "Warnings", "\n".join(f"• {w}" for w in confidence.warnings))

    for cell in table.columns[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Citations as numbered footnotes-style list (python-docx's footnote
    # API is anaemic; numbered list is the pragmatic choice).
    doc.add_heading("Citations", level=1)
    citations: Sequence[Any] = ctx.get("citations") or []
    if not citations:
        para = doc.add_paragraph(
            "No external citations were extracted from this run."
        )
        para.runs[0].italic = True
    else:
        for idx, c in enumerate(citations, start=1):
            line = doc.add_paragraph(style="List Number")
            label = c.title if hasattr(c, "title") else str(c)
            url = getattr(c, "url", None)
            file_path = getattr(c, "file_path", None)
            line.add_run(label).bold = False
            if url:
                line.add_run(f"  ({url})").italic = True
            elif file_path:
                line.add_run(f"  ({file_path})").italic = True
            section = getattr(c, "used_in_section", None)
            if section:
                tail = line.add_run(f"   — used in {section}")
                tail.italic = True
                tail.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            del idx  # numbering is implicit via the List Number style.

    # Appendix transcript.
    doc.add_heading("Appendix · Debate transcript", level=1)
    transcript = ctx.get("transcript_lines") or []
    if not transcript:
        doc.add_paragraph("(no transcript captured)").runs[0].italic = True
    else:
        for line in transcript:
            doc.add_paragraph(_strip_md(line), style="List Bullet")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _section(
    doc: Any,
    heading: str,
    body: Any,
    *,
    default: str | None = None,
    image: str | None = None,
    image_dir: str = "",
) -> None:
    doc.add_heading(heading, level=1)
    _embed_image(doc, image, image_dir, width_in=5.5)
    text = (body or default or "").strip()
    if not text:
        para = doc.add_paragraph("(no content)")
        para.runs[0].italic = True
        return
    for paragraph in text.split("\n\n"):
        doc.add_paragraph(paragraph.strip())


def _embed_image(
    doc: Any,
    rel_path: str | None,
    image_dir: str,
    *,
    width_in: float,
) -> None:
    """Inline a generated PNG into the document, if both inputs resolve.

    Silently no-ops when the image dir or relative path is missing —
    image generation is best-effort, so a per-section render miss
    must never break the report.
    """
    if not rel_path or not image_dir:
        return
    try:
        from pathlib import Path

        from docx.shared import Inches

        full = Path(image_dir) / rel_path
        if not full.exists():
            return
        doc.add_picture(str(full), width=Inches(width_in))
    except Exception:  # noqa: BLE001 — never crash a report on an image embed
        return


def _row(table: Any, key: str, value: str) -> None:
    cells = table.add_row().cells
    cells[0].text = key
    cells[1].text = value


def _strip_md(line: str) -> str:
    """Convert a Markdown bullet line to plain text for DOCX bullets."""
    s = line.lstrip("- ").strip()
    return s.replace("**", "")
